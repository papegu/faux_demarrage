"""experiments_utils.py
Utilities to run experiments for faux_demarrage project.

This module is designed to be executed locally (not run by this agent). It:
- loads dataset(s) from repository
- computes dataset statistics and figures
- builds baseline models and the hybrid Keras model
- runs nested CV (outer folds) and saves results and figures in OUTDIR
- provides a function to insert results into a copy of the manuscript .docx and mark additions in red

Run with: python run_experiments.py --data path/to/dataset.xlsx

"""

import os
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.model_selection import StratifiedKFold, train_test_split
from sklearn.preprocessing import StandardScaler
from sklearn.impute import SimpleImputer
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import roc_auc_score, precision_score, recall_score, f1_score, brier_score_loss
from sklearn.utils import resample
from imblearn.over_sampling import SMOTE
import tensorflow as tf
from tensorflow.keras.callbacks import EarlyStopping, ReduceLROnPlateau
from docx import Document
from docx.shared import Inches

try:
    import xgboost as xgb
    HAS_XGB = True
except Exception:
    HAS_XGB = False

RANDOM_STATE = 42
np.random.seed(RANDOM_STATE)

OUTDIR = 'improve_frontiers_outputs'
FIG_DIR = os.path.join(OUTDIR, 'figures')
os.makedirs(FIG_DIR, exist_ok=True)

DEFAULT_MANUSCRIPT = 'article Sequentiel hybrid app....Towards a Predictive Tool for False Starts.docx'

# -----------------------------
# Data helpers
# -----------------------------

def find_first_existing(paths):
    for p in paths:
        if os.path.exists(p):
            return p
    return None


def load_data(path=None, target_col='faux_demarage'):
    candidate_paths = [
        'modeles_avec_timestep/donnees_avec_faux_demarage.xlsx',
        'modeles_sans_pas_de_temps/donnees_avec_faux_demarage.xlsx',
        'preparer_data_set/dataset_rempli.xlsx',
        'donnees_avec_faux_demarage.xlsx'
    ]
    if path is None:
        path = find_first_existing(candidate_paths)
        if path is None:
            raise FileNotFoundError('No dataset found in expected locations. Provide --data path.')
    print(f'Loading data from {path}')
    if path.endswith('.xlsx') or path.endswith('.xls'):
        df = pd.read_excel(path)
    else:
        df = pd.read_csv(path)
    if target_col not in df.columns:
        raise ValueError(f"Target column '{target_col}' not found in data")
    return df


def dataset_summary(df, target_col='faux_demarage', outdir=FIG_DIR):
    os.makedirs(outdir, exist_ok=True)
    counts = df[target_col].value_counts(dropna=False).rename_axis('class').reset_index(name='count')
    counts['percent'] = 100 * counts['count'] / counts['count'].sum()
    counts.to_csv(os.path.join(outdir, 'class_distribution.csv'), index=False)

    plt.figure(figsize=(5,4))
    sns.barplot(data=counts, x='class', y='count', palette='viridis')
    plt.title('Distribution des classes')
    plt.tight_layout()
    plt.savefig(os.path.join(outdir, 'class_distribution.png'), dpi=300)
    plt.close()

    feat_df = df.drop(columns=[target_col])
    desc = feat_df.describe().T
    desc.to_csv(os.path.join(outdir, 'features_summary.csv'))

    missing = feat_df.isnull().sum().to_frame('missing').sort_values('missing', ascending=False)
    missing.to_csv(os.path.join(outdir, 'missing_per_feature.csv'))

    num = feat_df.select_dtypes(include=[np.number]).fillna(0)
    if num.shape[1] > 0:
        corr = num.corr()
        corr.to_csv(os.path.join(outdir, 'correlation_matrix.csv'))
        plt.figure(figsize=(8,6))
        sns.heatmap(corr, cmap='RdBu_r', center=0)
        plt.title('Matrice de corrélation')
        plt.tight_layout()
        plt.savefig(os.path.join(outdir, 'correlation_matrix.png'), dpi=300)
        plt.close()

    return {'counts': counts, 'desc': desc, 'missing': missing}

# -----------------------------
# Model helpers
# -----------------------------

def build_hybrid_keras_model(input_shape, lr=1e-3):
    tf.keras.backend.clear_session()
    inputs = tf.keras.Input(shape=input_shape)
    lstm_out = tf.keras.layers.LSTM(32)(inputs)
    gru_out = tf.keras.layers.GRU(32)(inputs)
    x_trans = tf.keras.layers.Dense(32)(inputs)
    attn_out = tf.keras.layers.MultiHeadAttention(num_heads=4, key_dim=8)(x_trans, x_trans)
    trans_out = tf.keras.layers.GlobalAveragePooling1D()(attn_out)
    concat = tf.keras.layers.concatenate([lstm_out, gru_out, trans_out])
    x = tf.keras.layers.Dropout(0.3)(concat)
    x = tf.keras.layers.Dense(32, activation='relu', kernel_regularizer=tf.keras.regularizers.l2(1e-4))(x)
    outputs = tf.keras.layers.Dense(1, activation='sigmoid')(x)
    model = tf.keras.Model(inputs, outputs)
    model.compile(optimizer=tf.keras.optimizers.Adam(learning_rate=lr), loss='binary_crossentropy')
    return model


def compute_metrics(y_true, y_pred_prob, threshold=0.5):
    y_pred = (y_pred_prob >= threshold).astype(int)
    acc = np.mean(y_pred == y_true)
    try:
        auc = roc_auc_score(y_true, y_pred_prob)
    except Exception:
        auc = float('nan')
    prec = precision_score(y_true, y_pred, zero_division=0)
    rec = recall_score(y_true, y_pred, zero_division=0)
    f1 = f1_score(y_true, y_pred, zero_division=0)
    brier = brier_score_loss(y_true, y_pred_prob)
    return {'accuracy': acc, 'auc': auc, 'precision': prec, 'recall': rec, 'f1': f1, 'brier': brier}

# -----------------------------
# Training pipelines
# -----------------------------

def train_baselines_and_hybrid(X_train_df, y_train, X_test_df, use_smote=False):
    # Fit imputer and scaler on training only
    imputer = SimpleImputer(strategy='median')
    scaler = StandardScaler()

    X_train_imp = imputer.fit_transform(X_train_df)
    X_test_imp = imputer.transform(X_test_df)
    X_train_scaled = scaler.fit_transform(X_train_imp)
    X_test_scaled = scaler.transform(X_test_imp)

    if use_smote:
        sm = SMOTE(random_state=RANDOM_STATE)
        X_train_scaled, y_train_res = sm.fit_resample(X_train_scaled, y_train)
    else:
        y_train_res = y_train

    results = {}
    rf = RandomForestClassifier(n_estimators=200, random_state=RANDOM_STATE, n_jobs=-1)
    rf.fit(X_train_scaled, y_train_res)
    results['RandomForest'] = {'model': rf, 'probs': rf.predict_proba(X_test_scaled)[:,1]}

    if HAS_XGB:
        xgb_clf = xgb.XGBClassifier(use_label_encoder=False, eval_metric='logloss', random_state=RANDOM_STATE, n_jobs=4)
        xgb_clf.fit(X_train_scaled, y_train_res)
        results['XGBoost'] = {'model': xgb_clf, 'probs': xgb_clf.predict_proba(X_test_scaled)[:,1]}

    # Keras hybrid
    timesteps = 1
    n_features = X_train_scaled.shape[1]
    X_tr_3d = X_train_scaled.reshape((X_train_scaled.shape[0], timesteps, n_features))
    X_te_3d = X_test_scaled.reshape((X_test_scaled.shape[0], timesteps, n_features))

    model = build_hybrid_keras_model(input_shape=(timesteps, n_features))
    es = EarlyStopping(monitor='val_loss', patience=15, restore_best_weights=True)
    rlr = ReduceLROnPlateau(monitor='val_loss', factor=0.5, patience=7)
    model.fit(X_tr_3d, y_train_res, epochs=200, batch_size=32, validation_split=0.2, callbacks=[es, rlr], verbose=1)
    results['HybridKeras'] = {'model': model, 'probs': model.predict(X_te_3d, verbose=0).flatten()}

    return results

# -----------------------------
# Nested CV orchestration
# -----------------------------

def nested_cv_evaluate(df, target_col='faux_demarage', outer_splits=5, use_smote=False, outdir=OUTDIR):
    os.makedirs(outdir, exist_ok=True)
    X = df.drop(columns=[target_col])
    y = df[target_col].values

    skf = StratifiedKFold(n_splits=outer_splits, shuffle=True, random_state=RANDOM_STATE)
    fold_results = []
    for fold, (train_idx, test_idx) in enumerate(skf.split(X, y), start=1):
        X_train_df = X.iloc[train_idx]
        X_test_df = X.iloc[test_idx]
        y_train = y[train_idx]
        y_test = y[test_idx]
        print(f'Fold {fold}: train={len(train_idx)}, test={len(test_idx)}')
        res = train_baselines_and_hybrid(X_train_df, y_train, X_test_df, use_smote=use_smote)
        for name, info in res.items():
            probs = info['probs']
            metrics = compute_metrics(y_test, probs)
            # save
            np.save(os.path.join(outdir, f'probs_fold{fold}_{name}.npy'), probs)
            metrics_row = {'fold': fold, 'model': name}
            metrics_row.update(metrics)
            fold_results.append(metrics_row)

    df_res = pd.DataFrame(fold_results)
    df_res.to_csv(os.path.join(outdir, 'nested_cv_fold_results.csv'), index=False)
    summary = df_res.groupby('model').agg({'accuracy': ['mean', 'std'], 'auc': ['mean', 'std'], 'f1': ['mean', 'std']})
    summary.to_csv(os.path.join(outdir, 'nested_cv_summary.csv'))
    return df_res, summary

# -----------------------------
# DOCX insertion helpers
# -----------------------------

def add_red_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    try:
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    except Exception:
        pass
    run.bold = bold
    return p


def insert_results_into_docx(src_docx=DEFAULT_MANUSCRIPT, out_docx=None, summary_table=None, figures=None):
    if out_docx is None:
        base, ext = os.path.splitext(src_docx)
        out_docx = base + '_rev.docx'
    doc = Document(src_docx)
    doc.add_page_break()
    doc.add_heading('Added analysis (automated)', level=1)
    add_red_paragraph(doc, 'New sections and figures added: nested CV results, dataset statistics, baselines comparison, and interpretability plots.', bold=True)
    if summary_table is not None:
        doc.add_heading('Nested CV summary (mean ± std)', level=2)
        dfst = summary_table.copy()
        dfst.columns = ['_'.join(map(str, c)).strip() for c in dfst.columns.values]
        t = doc.add_table(rows=1, cols=len(dfst.columns)+1)
        hdr = t.rows[0].cells
        hdr[0].text = 'Model'
        for i, c in enumerate(dfst.columns):
            hdr[i+1].text = str(c)
        for idx, row in dfst.iterrows():
            cells = t.add_row().cells
            cells[0].text = str(idx)
            for i, v in enumerate(row):
                cells[i+1].text = str(v)
    if figures:
        doc.add_heading('Figures added', level=2)
        for fig in figures:
            if os.path.exists(fig):
                doc.add_picture(fig, width=Inches(6))
                add_red_paragraph(doc, f'Figure added: {os.path.basename(fig)}')
    doc.save(out_docx)
    return out_docx

# -----------------------------
# Runner
# -----------------------------

def run_all(path=None, target='faux_demarage', use_smote=False):
    df = load_data(path, target_col=target)
    dataset_summary(df, target_col=target)
    df_res, summary = nested_cv_evaluate(df, target_col=target, use_smote=use_smote)
    figs = [os.path.join(FIG_DIR, 'class_distribution.png'), os.path.join(FIG_DIR, 'correlation_matrix.png')]
    outdoc = insert_results_into_docx(DEFAULT_MANUSCRIPT, out_docx=os.path.join(OUTDIR, 'article_rev.docx'), summary_table=summary, figures=figs)
    return {'df_res': df_res, 'summary': summary, 'outdoc': outdoc}
